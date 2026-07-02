import json
from pathlib import Path

from docx import Document

from src.config import Settings
from src.services.block_builder import build_blocks_json


def test_build_pdf_blocks_from_document_units(tmp_path: Path) -> None:
    settings = make_settings(tmp_path)
    write_document(
        settings,
        "file-uuid",
        {
            "texts": [
                {
                    "text": "第一章 协议总则",
                    "prov": [
                        {
                            "page_no": 1,
                            "line_index": 0,
                            "bbox": {"l": 10, "t": 20, "r": 210, "b": 40},
                        }
                    ],
                },
                {
                    "text": "本章正文",
                    "prov": [
                        {
                            "page_no": 1,
                            "line_index": 1,
                            "bbox": {"l": 10, "t": 50, "r": 180, "b": 70},
                        }
                    ],
                },
            ],
            "tables": [
                {
                    "text": "| 名称 | 内容 |",
                    "page_no": 2,
                    "line_index": 0,
                    "bbox": {"l": 20, "t": 30, "r": 220, "b": 80},
                }
            ],
        },
    )

    payload = build_blocks_json("file-uuid", tmp_path / "source.pdf", ".pdf", settings)

    assert payload == {
        "total_pages": 2,
        "blocks": [
            {
                "no": 1,
                "page": 1,
                "text": "第一章 协议总则",
                "kind": "chapter",
                "rank": 0,
                "bbox": {"x": 10, "y": 20, "width": 200, "height": 20},
            },
            {
                "no": 2,
                "page": 1,
                "text": "本章正文",
                "kind": None,
                "rank": None,
                "bbox": {"x": 10, "y": 50, "width": 170, "height": 20},
            },
            {
                "no": 3,
                "page": 2,
                "text": "| 名称 | 内容 |",
                "kind": None,
                "rank": None,
                "bbox": {"x": 20, "y": 30, "width": 200, "height": 50},
            },
        ],
    }


def test_build_docx_blocks_matches_markdown_lines_to_document_units(tmp_path: Path) -> None:
    settings = make_settings(tmp_path)
    source_path = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("第一章 协议总则")
    document.add_paragraph("本章正文")
    document.save(source_path)
    write_document(
        settings,
        "file-uuid",
        {
            "texts": [
                {
                    "text": "第一章 协议总则",
                    "prov": [
                        {
                            "page_no": 1,
                            "line_index": 0,
                            "bbox": {"l": 10, "t": 20, "r": 210, "b": 40},
                        }
                    ],
                },
                {
                    "text": "本章正文",
                    "prov": [
                        {
                            "page_no": 1,
                            "line_index": 1,
                            "bbox": {"l": 10, "t": 50, "r": 180, "b": 70},
                        }
                    ],
                },
            ],
            "tables": [],
        },
    )

    payload = build_blocks_json("file-uuid", source_path, ".docx", settings)

    assert payload["total_pages"] == 1
    assert payload["blocks"][0] == {
        "no": 1,
        "page": 1,
        "text": "第一章 协议总则",
        "kind": "chapter",
        "rank": 0,
        "bbox": {"x": 10, "y": 20, "width": 200, "height": 50},
    }
    assert payload["blocks"][1] == {
        "no": 2,
        "page": 1,
        "text": "本章正文",
        "kind": None,
        "rank": None,
        "bbox": {"x": 10, "y": 50, "width": 170, "height": 20},
    }


def make_settings(tmp_path: Path) -> Settings:
    return Settings(
        upload_dir=tmp_path / "uploads",
        parsing_dir=tmp_path / "parsing",
    )


def write_document(settings: Settings, file_id: str, document: dict) -> None:
    document_dir = settings.parsing_root / file_id
    document_dir.mkdir(parents=True, exist_ok=True)
    (document_dir / "document.json").write_text(
        json.dumps(document, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
