from docx import Document

from src.services.python_docx_contract_parser import (
    parse_docx_contract_to_json_data,
    parse_number_token,
)


def test_parse_number_token_assigns_contract_ranks():
    assert parse_number_token("第一部分 专用条款").rank == 0
    assert parse_number_token("附件一 授权委托书").rank == 5
    assert parse_number_token("一、工程概况").rank == 10
    assert parse_number_token("第一条 共同约定").rank == 10
    assert parse_number_token("1、工程名称").rank == 20
    assert parse_number_token("1. 乙方义务").rank == 20
    assert parse_number_token("1.1 本合同为固定单价合同").rank == 21
    assert parse_number_token("1.1.1 合规开展行业竞争").rank == 22
    assert parse_number_token("（1）资料").rank == 40


def test_parse_number_token_prefers_literal_label_when_auto_label_conflicts():
    token = parse_number_token("1.1 手工编号", auto_label="1、")

    assert token.raw_label == "1.1"
    assert token.auto_label == "1、"
    assert token.literal_label == "1.1"
    assert token.warnings == ["自动编号与文本编号同时存在"]


def test_parse_docx_contract_to_json_data_builds_structure_and_tables(tmp_path):
    document = Document()
    document.add_paragraph("第一部分 专用条款")
    document.add_paragraph("一、工程概况")
    document.add_paragraph("1、工程名称")
    document.add_paragraph("北京项目")
    document.add_paragraph("1.1 本合同为固定单价合同")
    document.add_paragraph("1.3 跳过一个编号")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "1、表格内部序号"
    table.cell(0, 1).text = "货物名称"
    document.add_paragraph("附件一 授权委托书")
    document.add_paragraph("附件二 项目部合规权限告知书")
    document.add_paragraph("附件一：")

    docx_path = tmp_path / "contract.docx"
    document.save(docx_path)

    payload = parse_docx_contract_to_json_data(docx_path, doc_id="file-uuid")

    root = payload["contract_structure"]
    part = root["children"][0]
    chinese = part["children"][0]
    arabic = chinese["children"][0]
    first_dotted = arabic["children"][0]
    second_dotted = arabic["children"][1]

    assert payload["schema_name"] == "PythonDocxContractDocument"
    assert payload["name"] == "file-uuid"
    assert len(payload["texts"]) == 9
    assert part["label"] == "第一部分"
    assert chinese["label"] == "一、"
    assert arabic["label"] == "1、"
    assert arabic["content"][0]["text"] == "北京项目"
    assert first_dotted["label"] == "1.1"
    assert second_dotted["label"] == "1.3"
    assert "同级编号不连续" in second_dotted["warnings"][0]
    assert second_dotted["tables"][0]["rows"] == [["1、表格内部序号", "货物名称"]]
    attachments = [child for child in part["children"] if child["kind"] == "attachment"]
    assert len(attachments) == 1
    assert attachments[0]["label"] == "附件一"
    assert attachments[0]["title"] == ""
    assert all(item["text"] != "1、表格内部序号" for item in payload["texts"])


def test_parse_docx_contract_to_json_data_skips_toc_and_adds_preamble(tmp_path):
    document = Document()
    document.add_paragraph("大数据应用系统及数据灾备系统平台开发项目服务合同")
    document.add_paragraph("委托方（甲方）：A公司")
    document.add_paragraph("受托方（乙方）：B公司")
    document.add_paragraph("目 录")
    document.add_paragraph("定义")
    document.add_paragraph("第一章 项目名称和开发依据")
    document.add_paragraph("第二章 合同价款及支付方式")
    document.add_paragraph("附件")
    document.add_paragraph("附件一 服务清单")
    document.add_paragraph("附加五 安全保密协议")
    for _ in range(4):
        document.add_paragraph("")
    document.add_paragraph("定义")
    document.add_paragraph("“本工程”或“本项目”：指大数据应用系统开发项目。")
    document.add_paragraph("第一章 项目名称和开发依据")
    document.add_paragraph("本合同的委托开发项目名称为：大数据应用系统。")

    docx_path = tmp_path / "contract.docx"
    document.save(docx_path)

    payload = parse_docx_contract_to_json_data(docx_path, doc_id="file-uuid")

    root = payload["contract_structure"]
    children = root["children"]
    preamble = children[0]
    definition = children[1]
    chapter = children[2]

    assert preamble["kind"] == "preamble"
    assert preamble["title"] == "合同首部"
    assert [item["text"] for item in preamble["content"]] == [
        "大数据应用系统及数据灾备系统平台开发项目服务合同",
        "委托方（甲方）：A公司",
        "受托方（乙方）：B公司",
    ]
    assert definition["kind"] == "heading"
    assert definition["title"] == "定义"
    assert definition["content"][0]["text"].startswith("“本工程”")
    assert chapter["label"] == "第一章"
    assert chapter["title"] == "项目名称和开发依据"
    assert all(child["title"] != "合同价款及支付方式" for child in children[:3])
