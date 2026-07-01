import json
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from src.config import Settings
from src.services import document_parser
from src.services.document_parser import ParsingError


def test_build_document_converter_reuses_cached_instance(monkeypatch):
    created = []

    def create_document_converter():
        converter = object()
        created.append(converter)
        return converter

    monkeypatch.setattr(document_parser, "_DOCUMENT_CONVERTER", None)
    monkeypatch.setattr(
        document_parser,
        "create_document_converter",
        create_document_converter,
    )

    first = document_parser.build_document_converter()
    second = document_parser.build_document_converter()

    assert first is second
    assert created == [first]


def test_warmup_document_converter_initializes_pdf_pipeline(monkeypatch):
    initialized_formats = []

    class FakeConverter:
        def initialize_pipeline(self, doc_format):
            initialized_formats.append(doc_format)

    class FakeInputFormat:
        PDF = "pdf"

    monkeypatch.setattr(
        document_parser,
        "build_document_converter",
        lambda: FakeConverter(),
    )
    monkeypatch.setattr(
        document_parser,
        "load_docling_types",
        lambda: (None, None, FakeInputFormat, None),
    )

    document_parser.warmup_document_converter()

    assert initialized_formats == [FakeInputFormat.PDF]


def test_pdf_parsing_writes_json_to_uuid_directory(tmp_path, monkeypatch):
    source_path = tmp_path / "source.pdf"
    source_path.write_bytes(b"%PDF-1.4\n%%EOF\n")
    settings = make_settings(tmp_path)

    def fake_convert_to_json_data(input_path: Path):
        assert input_path == source_path
        return {
            "texts": [
                {
                    "self_ref": "#/texts/0",
                    "text": "第一部分 专用条款",
                    "prov": [{"page_no": 1}],
                }
            ]
        }

    monkeypatch.setattr(
        document_parser,
        "convert_to_json_data",
        fake_convert_to_json_data,
    )

    artifacts = document_parser.parse_uploaded_document(
        source_path,
        "file-uuid",
        ".pdf",
        settings,
    )

    assert artifacts.directory == settings.parsing_root / "file-uuid"
    assert artifacts.json_path == artifacts.directory / "document.json"
    assert artifacts.pdf_path is None
    payload = artifacts.json_path.read_text(encoding="utf-8")
    assert '"text": "第一部分 专用条款"' in payload
    assert '"contract_structure"' in payload


def test_pdf_parsing_can_use_pymupdf4llm_engine(tmp_path, monkeypatch):
    source_path = tmp_path / "source.pdf"
    source_path.write_bytes(b"%PDF-1.4\n%%EOF\n")
    settings = make_settings(tmp_path, pdf_parser_engine="pymupdf4llm")

    def fail_convert_to_json_data(input_path: Path):
        raise AssertionError("pymupdf4llm engine should not use Docling")

    def fake_convert_pdf_to_pymupdf4llm_json_data(input_path: Path, file_id: str):
        assert input_path == source_path
        assert file_id == "file-uuid"
        return {
            "schema_name": "PyMuPDF4LLMContractDocument",
            "texts": [{"self_ref": "#/texts/0", "text": "合同"}],
        }

    monkeypatch.setattr(document_parser, "convert_to_json_data", fail_convert_to_json_data)
    monkeypatch.setattr(
        document_parser,
        "convert_pdf_to_pymupdf4llm_json_data",
        fake_convert_pdf_to_pymupdf4llm_json_data,
    )

    artifacts = document_parser.parse_uploaded_document(
        source_path,
        "file-uuid",
        ".pdf",
        settings,
    )

    payload = artifacts.json_path.read_text(encoding="utf-8")
    assert artifacts.pdf_path is None
    assert '"schema_name": "PyMuPDF4LLMContractDocument"' in payload


def test_pdf_parsing_rejects_unknown_pdf_engine(tmp_path):
    source_path = tmp_path / "source.pdf"
    source_path.write_bytes(b"%PDF-1.4\n%%EOF\n")
    settings = make_settings(tmp_path, pdf_parser_engine="bad-engine")

    with pytest.raises(ParsingError) as exc_info:
        document_parser.parse_uploaded_document(
            source_path,
            "file-uuid",
            ".pdf",
            settings,
        )

    assert exc_info.value.code == "UNSUPPORTED_PDF_PARSER_ENGINE"


def test_docx_to_pdf_disables_libreoffice_bookmark_export(tmp_path, monkeypatch):
    input_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.pdf"
    input_path.write_bytes(b"docx")
    commands = []

    def fake_run(command, text, capture_output, check):
        commands.append(command)
        out_dir = Path(command[command.index("--outdir") + 1])
        source = Path(command[-1])
        (out_dir / f"{source.stem}.pdf").write_bytes(b"%PDF-1.4\n%%EOF\n")
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr(
        document_parser,
        "find_libreoffice_executable",
        lambda: "/usr/bin/libreoffice",
    )
    monkeypatch.setattr(document_parser.subprocess, "run", fake_run)

    result_path = document_parser.convert_docx_to_pdf(input_path, output_path)

    filter_arg = commands[0][commands[0].index("--convert-to") + 1]
    prefix = "pdf:writer_pdf_Export:"
    options = json.loads(filter_arg.removeprefix(prefix))

    assert result_path == output_path
    assert output_path.read_bytes() == b"%PDF-1.4\n%%EOF\n"
    assert filter_arg.startswith(prefix)
    assert options["ExportBookmarks"] == {"type": "boolean", "value": "false"}
    assert options["InitialView"] == {"type": "long", "value": "0"}


def test_docx_parsing_saves_converted_pdf_then_json(tmp_path, monkeypatch):
    source_path = tmp_path / "source.docx"
    source_path.write_bytes(b"docx")
    settings = make_settings(tmp_path)
    converted_inputs = []
    cleaned_inputs = []

    def fake_remove_docx_headers_footers(input_path: Path, output_path: Path):
        assert input_path == source_path
        assert output_path == settings.parsing_root / "file-uuid" / "document.no-header-footer.docx"
        output_path.write_bytes(b"clean docx")
        cleaned_inputs.append(output_path)
        return output_path

    def fake_convert_docx_to_pdf(input_path: Path, output_path: Path):
        assert input_path == settings.parsing_root / "file-uuid" / "document.no-header-footer.docx"
        output_path.write_bytes(b"%PDF-1.4\n%%EOF\n")
        return output_path

    def fake_convert_to_json_data(input_path: Path):
        converted_inputs.append(input_path)
        return {"texts": ["合同"]}

    monkeypatch.setattr(
        document_parser,
        "remove_docx_headers_footers",
        fake_remove_docx_headers_footers,
    )
    monkeypatch.setattr(
        document_parser,
        "convert_docx_to_pdf",
        fake_convert_docx_to_pdf,
    )
    monkeypatch.setattr(
        document_parser,
        "convert_to_json_data",
        fake_convert_to_json_data,
    )

    artifacts = document_parser.parse_uploaded_document(
        source_path,
        "file-uuid",
        ".docx",
        settings,
    )

    assert artifacts.directory == settings.parsing_root / "file-uuid"
    assert artifacts.pdf_path == artifacts.directory / "document.pdf"
    assert artifacts.pdf_path.read_bytes() == b"%PDF-1.4\n%%EOF\n"
    assert artifacts.json_path.is_file()
    assert cleaned_inputs == [artifacts.directory / "document.no-header-footer.docx"]
    assert converted_inputs == [artifacts.pdf_path]


def test_docx_parsing_can_use_pymupdf4llm_engine_after_pdf_conversion(
    tmp_path,
    monkeypatch,
):
    source_path = tmp_path / "source.docx"
    source_path.write_bytes(b"docx")
    settings = make_settings(tmp_path, docx_parser_engine="pymupdf4llm")
    parsed_inputs = []
    cleaned_inputs = []

    def fake_remove_docx_headers_footers(input_path: Path, output_path: Path):
        assert input_path == source_path
        assert output_path == settings.parsing_root / "file-uuid" / "document.no-header-footer.docx"
        output_path.write_bytes(b"clean docx")
        cleaned_inputs.append(output_path)
        return output_path

    def fake_convert_docx_to_pdf(input_path: Path, output_path: Path):
        assert input_path == settings.parsing_root / "file-uuid" / "document.no-header-footer.docx"
        output_path.write_bytes(b"%PDF-1.4\n%%EOF\n")
        return output_path

    def fail_convert_to_json_data(input_path: Path):
        raise AssertionError("pymupdf4llm DOCX engine should not use Docling")

    def fake_convert_pdf_to_pymupdf4llm_json_data(input_path: Path, file_id: str):
        parsed_inputs.append(input_path)
        assert file_id == "file-uuid"
        return {
            "schema_name": "PyMuPDF4LLMContractDocument",
            "texts": [{"self_ref": "#/texts/0", "text": "合同"}],
        }

    monkeypatch.setattr(
        document_parser,
        "remove_docx_headers_footers",
        fake_remove_docx_headers_footers,
    )
    monkeypatch.setattr(
        document_parser,
        "convert_docx_to_pdf",
        fake_convert_docx_to_pdf,
    )
    monkeypatch.setattr(document_parser, "convert_to_json_data", fail_convert_to_json_data)
    monkeypatch.setattr(
        document_parser,
        "convert_pdf_to_pymupdf4llm_json_data",
        fake_convert_pdf_to_pymupdf4llm_json_data,
    )

    artifacts = document_parser.parse_uploaded_document(
        source_path,
        "file-uuid",
        ".docx",
        settings,
    )

    payload = artifacts.json_path.read_text(encoding="utf-8")
    assert artifacts.pdf_path == artifacts.directory / "document.pdf"
    assert artifacts.pdf_path.read_bytes() == b"%PDF-1.4\n%%EOF\n"
    assert cleaned_inputs == [artifacts.directory / "document.no-header-footer.docx"]
    assert parsed_inputs == [artifacts.pdf_path]
    assert '"schema_name": "PyMuPDF4LLMContractDocument"' in payload


def test_docx_parsing_can_use_python_docx_engine(tmp_path, monkeypatch):
    source_path = tmp_path / "source.docx"
    source_path.write_bytes(b"docx")
    settings = make_settings(tmp_path, docx_parser_engine="python-docx")

    def fail_convert_docx_to_pdf(input_path: Path, output_path: Path):
        raise AssertionError("python-docx engine should not convert DOCX to PDF")

    def fake_convert_docx_to_python_docx_json_data(input_path: Path, file_id: str):
        assert input_path == source_path
        assert file_id == "file-uuid"
        return {
            "schema_name": "PythonDocxContractDocument",
            "texts": [{"self_ref": "#/texts/0", "text": "合同"}],
        }

    monkeypatch.setattr(
        document_parser,
        "convert_docx_to_pdf",
        fail_convert_docx_to_pdf,
    )
    monkeypatch.setattr(
        document_parser,
        "convert_docx_to_python_docx_json_data",
        fake_convert_docx_to_python_docx_json_data,
    )

    artifacts = document_parser.parse_uploaded_document(
        source_path,
        "file-uuid",
        ".docx",
        settings,
    )

    payload = artifacts.json_path.read_text(encoding="utf-8")
    assert artifacts.pdf_path is None
    assert '"schema_name": "PythonDocxContractDocument"' in payload


def test_remove_docx_headers_footers_preserves_body_and_clears_story_parts(tmp_path):
    input_path = tmp_path / "source.docx"
    output_path = tmp_path / "clean.docx"

    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "页眉内容"
    section.footer.paragraphs[0].text = "页脚内容"
    section.header.add_table(rows=1, cols=1, width=1).cell(0, 0).text = "页眉表格"
    document.add_paragraph("正文内容")
    document.save(input_path)

    result_path = document_parser.remove_docx_headers_footers(input_path, output_path)

    cleaned = Document(result_path)
    cleaned_section = cleaned.sections[0]

    assert result_path == output_path
    assert [paragraph.text for paragraph in cleaned.paragraphs] == ["正文内容"]
    assert [paragraph.text for paragraph in cleaned_section.header.paragraphs] == [""]
    assert [paragraph.text for paragraph in cleaned_section.footer.paragraphs] == [""]
    assert cleaned_section.header.tables == []


def test_docx_parsing_rejects_unknown_docx_engine(tmp_path):
    source_path = tmp_path / "source.docx"
    source_path.write_bytes(b"docx")
    settings = make_settings(tmp_path, docx_parser_engine="bad-engine")

    with pytest.raises(ParsingError) as exc_info:
        document_parser.parse_uploaded_document(
            source_path,
            "file-uuid",
            ".docx",
            settings,
        )

    assert exc_info.value.code == "UNSUPPORTED_DOCX_PARSER_ENGINE"


def test_parse_failure_removes_partial_directory(tmp_path, monkeypatch):
    source_path = tmp_path / "source.pdf"
    source_path.write_bytes(b"%PDF-1.4\n%%EOF\n")
    settings = make_settings(tmp_path)

    def fail_convert_to_json_data(input_path: Path):
        partial_path = settings.parsing_root / "file-uuid" / "partial.txt"
        partial_path.write_text("partial", encoding="utf-8")
        raise ValueError("parse failed")

    monkeypatch.setattr(
        document_parser,
        "convert_to_json_data",
        fail_convert_to_json_data,
    )

    with pytest.raises(ParsingError) as exc_info:
        document_parser.parse_uploaded_document(
            source_path,
            "file-uuid",
            ".pdf",
            settings,
        )

    assert exc_info.value.code == "DOCUMENT_PARSING_ERROR"
    assert not (settings.parsing_root / "file-uuid").exists()


def test_unexpected_storage_failure_is_wrapped_and_cleaned(tmp_path, monkeypatch):
    source_path = tmp_path / "source.pdf"
    source_path.write_bytes(b"%PDF-1.4\n%%EOF\n")
    settings = make_settings(tmp_path)

    monkeypatch.setattr(
        document_parser,
        "convert_to_json_data",
        lambda input_path: {"texts": ["合同"]},
    )

    def fail_write_json(path: Path, data):
        partial_path = path.parent / "document.json.part"
        partial_path.write_text("partial", encoding="utf-8")
        raise OSError("write failed")

    monkeypatch.setattr(document_parser, "write_json_atomically", fail_write_json)

    with pytest.raises(ParsingError) as exc_info:
        document_parser.parse_uploaded_document(
            source_path,
            "file-uuid",
            ".pdf",
            settings,
        )

    assert exc_info.value.code == "DOCUMENT_PARSING_ERROR"
    assert exc_info.value.status_code == 500
    assert not (settings.parsing_root / "file-uuid").exists()


def make_settings(tmp_path, **kwargs):
    kwargs.setdefault("docx_parser_engine", "docling")
    kwargs.setdefault("pdf_parser_engine", "docling")
    return Settings(
        upload_dir=tmp_path / "uploads",
        parsing_dir=tmp_path / "parsing",
        max_upload_size_mb=50,
        **kwargs,
    )
